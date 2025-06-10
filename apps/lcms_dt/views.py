import re
import traceback
import datetime

from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status

from docx import Document
import xlrd3 as xlrd
from .models import Lcms_UserFiles


class LcmsProcessAb(APIView):
    # AB液质
    def post(self, request):
        user = request.user
        # 判断是否登录
        if not user.is_authenticated:
            return Response(
                {"error": "Authentication required"},
                status=status.HTTP_401_UNAUTHORIZED
            )
        # 获取用户uuid
        uuid = user.uuid
        # 生成本次处理记录的id（系统唯一）
        process_id = ""
        current_time = datetime.datetime.now()
        current_time_str = str(current_time.year) + str(current_time.month) + str(current_time.day) + str(
            current_time.hour) + str(current_time.minute) + str(current_time.second)
        process_id = str(uuid) + current_time_str + "lcms"
        # 获取所有上传文件对象
        uploaded_files = request.FILES.getlist('files')
        if not uploaded_files:
            return Response({"status": "error", "message": "未检测到文件"}, status=400)
        try:
            # 处理结果
            processed_results = []
            # 遍历处理每个文件，得到每个文件的待处理数据集合
            for file_obj in uploaded_files:
                # 记录单个文件状态
                file_result = {
                    "file_name": file_obj.name,
                    "status": "success",
                    "compound_list": [],
                    "data": [],
                    "errors": []
                }
                try:
                    # 检验文件类型
                    if not file_obj.name.lower().endswith('.docx'):
                        raise ValueError("仅支持Word文件 (.docx)")

                    # 读取Doc内容
                    doc = Document(file_obj)
                    # 提取物质名称
                    file_compound = []
                    paras = doc.paragraphs  # 文档中段落的集合
                    for para in paras:
                        if para.text.strip(' \n') != '':
                            temp = re.search("Analyte: (.*? \(.*?\))", para.text)
                            file_compound.append(temp.group(1))
                        else:
                            continue
                    file_result["compound_list"] = file_compound
                    # 提取第一个表格
                    tables = doc.tables     # 当前文件中的所有表格对象
                    process_table = tables[1]
                    columns = process_table._column_count  # 列的长度
                    rows = len(process_table.rows)  # 行的长度
                    # 获取当前文件的表格表头（第一个表格的第一行）
                    head = []
                    for c in range(0, columns):
                        cell = process_table.cell(0, c)
                        txt = cell.text.replace('\n', '')
                        head.append(txt)
                    # 获取当前文件的样品名称列表（第一个表格的第一列）
                    sample_name = []
                    for c in range(1, rows):
                        cell = process_table.cell(c, 0)
                        sample_name.append(cell.text)

                    # 获取所有化合物对应的数据
                    # 举例：样品列表如有60个，则对应有60张表的样品详细数据，外加每一部分开始的两张表，则每一部分共62张表。
                    # 每个文件中有len(file_compound)个部分，每一部分有2+len(sample_name)张表
                    # 待处理的数据为每一部分的第二张表（索引为1）
                    file_data = []
                    for table_idx in range(1, len(tables), len(sample_name) + 2):
                        table_data = []
                        for row in tables[table_idx].rows:
                            row_data = [cell.text for cell in row.cells]
                            table_data.append(row_data)

                        file_data.append(table_data)

                    # 单个文件待处理数据集合
                    file_result["data"] = file_data

                    # print(len(file_compound))
                    # print(len(file_result["data"]))
                    # print("----")

                except Exception as e:
                    file_result["errors"].append({
                        "error": str(e)
                    })

                # 添加单个文件的待处理数据对象
                processed_results.append(file_result)

            # for result in processed_results:
            #     print(result["data"])
            #     print("----------")

            for file_obj in processed_results:
                for data_idx in range(len(file_obj["data"])):
                    single_result_head = file_obj["data"][data_idx][0]  # 获取表头
                    # 去除表头元素
                    file_obj["data"][data_idx] = file_obj["data"][data_idx][1:]
                    # 源数据
                    origin_data = file_obj["data"][data_idx]
                    # 转换后数据
                    transform_data = [
                        {single_result_head[i]: value for i, value in enumerate(item)}
                        for item in origin_data
                    ]
                    file_obj["data"][data_idx] = transform_data

            # 文件处理结束，将处理结果保存至数据库
            Lcms_UserFiles.objects.create(process_id=process_id, user=user, file_type="lcms-ab", single_file_json=processed_results)

        except Exception as e:
            return Response({
                "status": "error",
                "message": "服务器处理异常",
                "detail": str(e)
            }, status=500)

        return Response({
            "status": "completed",
            "total_files": len(processed_results),
            "single_results": processed_results
        })


class LcmsProcessAjl(APIView):
    # 安捷伦液质（6470）
    def post(self, request):
        user = request.user
        # 判断是否登录
        if not user.is_authenticated:
            return Response(
                {"error": "Authentication required"},
                status=status.HTTP_401_UNAUTHORIZED
            )
        # 获取用户uuid
        uuid = user.uuid
        # 生成本次处理记录的id（系统唯一）
        process_id = ""
        current_time = datetime.datetime.now()
        current_time_str = str(current_time.year) + str(current_time.month) + str(current_time.day) + str(
            current_time.hour) + str(current_time.minute) + str(current_time.second)
        process_id = str(uuid) + current_time_str + "lcms"
        # 获取所有上传文件对象
        uploaded_files = request.FILES.getlist('files')
        if not uploaded_files:
            return Response({"status": "error", "message": "未检测到文件"}, status=400)
        try:
            # 处理结果
            processed_results = []
            # 遍历处理每个文件，得到每个文件的待处理数据集合
            for file_obj in uploaded_files:
                # 记录单个文件状态
                file_result = {
                    "file_name": file_obj.name,
                    "status": "success",
                    "concentration_unit": "",
                    "data": [],
                    "errors": []
                }
                try:
                    # 检验文件类型
                    if not file_obj.name.lower().endswith(('.xls', '.xlsx')):
                        raise ValueError("仅支持Excel文件 (.xls/.xlsx)")

                    # 读取Excel内容
                    file_data = file_obj.read()
                    workbook = xlrd.open_workbook(file_contents=file_data)
                    sheets = workbook.sheet_names()
                    sheets = sheets[2:-1]

                    # 读取文件内容
                    file_data = []
                    for sheet in sheets:
                        process_sheet = workbook.sheet_by_name(sheet)
                        # 单独处理第一个Sheet表，提取固定数据（包括[数据文件，样品类型]）
                        if sheet == sheets[0]:
                            flag = False      # 标志是否提取区域
                            for row_idx in range(0, process_sheet.nrows):
                                flag_cell = process_sheet.cell(row_idx, 0)  # 每行列索引0的单元格值
                                if flag_cell.value == "":
                                    flag = False
                                # 提取区域临时中断
                                if flag and flag_cell.value == "数据文件":
                                    continue
                                # 提取区域开始
                                if flag_cell.value == "数据文件":
                                    # 列索引26为浓度单位
                                    unit = process_sheet.cell(row_idx + 1, 26).value
                                    file_result["concentration_unit"] = unit
                                    flag = True
                                # 提取固定数据
                                if flag:
                                    # 0：数据文件（充当样品名称），6：样品类型
                                    row_temp = [flag_cell.value, process_sheet.cell(row_idx, 6).value]
                                    file_data.append(row_temp)

                            # for data in file_data:
                            #     print(data)
                            # print("######")

                        # 获取每个化合物对应sheet中的浓度数据
                        file_concentration = []  # 当前sheet中的浓度数据
                        flag = False    # 标志是否提取区域
                        for row_idx in range(0, process_sheet.nrows):
                            flag_cell = process_sheet.cell(row_idx, 0)
                            if flag_cell.value == "":
                                flag = False
                            if flag and flag_cell.value == "数据文件":
                                continue
                            if flag_cell.value == "数据文件":
                                flag = True
                            if flag:
                                # 列索引为23的单元格为浓度值
                                file_concentration.append(process_sheet.cell(row_idx, 23).value)
                        file_concentration[0] = sheet
                        # 将浓度数据添加至最终处理结果列表
                        for data_idx in range(0, len(file_data)):
                            file_data[data_idx].append(file_concentration[data_idx])

                        # for data in file_data:
                        #     print(data)
                        # print("----")
                        # print(len(file_data))

                        # 单个文件待处理数据集合
                        file_result["data"] = file_data

                except Exception as e:
                    exception_traceback = traceback.format_exc()
                    print("异常堆栈信息：", exception_traceback)
                    file_result["errors"].append({
                        "error": str(e)
                    })

                # 添加单个文件的待处理数据对象
                processed_results.append(file_result)

            # for result in processed_results:
            #     print(result["data"])
            #     print("----------")

            for file_obj in processed_results:
                # 获取表头
                single_result_head = file_obj["data"][0]
                # 去除表头元素
                file_obj["data"] = file_obj["data"][1:]
                # 源数据
                origin_data = file_obj["data"]
                # 转换后数据
                transform_data = [
                    {single_result_head[i]: value for i, value in enumerate(item)}
                    for item in origin_data
                ]
                file_obj["data"] = transform_data

            # 文件处理结束，将处理结果保存至数据库
            Lcms_UserFiles.objects.create(process_id=process_id, user=user, file_type="lcms-ajl-6470", single_file_json=processed_results)

        except Exception as e:
            return Response({
                "status": "error",
                "message": "服务器处理异常",
                "detail": str(e)
            }, status=500)

        return Response({
            "status": "completed",
            "total_files": len(processed_results),
            "single_results": processed_results
        })
